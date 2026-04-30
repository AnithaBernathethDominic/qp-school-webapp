import io
import os
import re
from collections import Counter
from datetime import datetime

import fitz
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def extract_pdf_text(path_or_file):
    if isinstance(path_or_file, (str, os.PathLike)):
        doc = fitz.open(path_or_file)
    else:
        data = path_or_file.read()
        doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for i, page in enumerate(doc, start=1):
        pages.append((i, page.get_text("text")))
    doc.close()
    return pages


def parse_syllabus(pages):
    topic = None
    chapter = None
    rows = []
    topic_re = re.compile(r"Topic\s+(\d+)\s*[-–]\s*(.+)", re.I)
    chapter_re = re.compile(r"^(\d+\.\d+)\s+(.+)$")
    sub_re = re.compile(r"^(\d+\.\d+\.\d+)\s*(?:✓|✔)?\s*(.+)$")
    for page_no, text in pages:
        for raw in text.splitlines():
            line = raw.strip().replace("\uf0fc", "✓")
            if not line:
                continue
            m = topic_re.match(line)
            if m:
                topic = f"Topic {m.group(1)} - {m.group(2).strip()}"
                chapter = None
                continue
            m = chapter_re.match(line)
            if m:
                chapter = f"{m.group(1)} {m.group(2).strip()}"
                continue
            m = sub_re.match(line)
            if m and topic and chapter:
                code = m.group(1)
                name = m.group(2).replace("✓", "").strip()
                rows.append({"topic": topic, "chapter": chapter, "code": code, "subtopic": name})
    return rows


def parse_questions(pages):
    full = "\n".join(text for _, text in pages)
    start = re.search(r"(?:^|\n)1\.\s+", full)
    if start:
        full = full[start.start():]
    pattern = re.compile(r"(?:^|\n)(\d{1,3})\.\s+(.*?)(?=\n\d{1,3}\.\s+|\Z)", re.S)
    questions = []
    for m in pattern.finditer(full):
        qno = int(m.group(1))
        body = re.sub(r"Page\s+\d+\s+of\s+\d+", "", m.group(2), flags=re.I)
        body = re.sub(r"\s+", " ", body).strip()
        if body:
            questions.append({"question_no": qno, "question": body})
    return sorted({q["question_no"]: q for q in questions}.values(), key=lambda x: x["question_no"])


RULES = [
    ("1.4.4", ["spring", "extension", "force and extension"]),
    ("1.4.5", ["hooke", "spring constant"]),
    ("1.5.1", ["moment", "pivot", "metre rule", "door handle", "clockwise", "anticlockwise"]),
    ("1.5.2", ["balanced", "equilibrium", "tilted", "balance"]),
    ("1.5.3", ["centre of gravity", "center of gravity", "hanging", "nail"]),
    ("1.8.2", ["coal", "natural gas", "oil", "fossil", "biofuel", "greenhouse", "carbon capture", "sulphur", "fuels"]),
    ("1.8.3", ["hydroelectric", "tidal", "water"]),
    ("1.8.1", ["radiation from the sun", "solar", "sun"]),
    ("1.8.4", ["geothermal"]),
    ("1.8.5", ["nuclear fusion"]),
    ("1.7.6", ["power", "work done per unit time", "energy transferred per second", "per second", "time taken"]),
    ("1.7.5", ["work done", "force acting over a distance", "braking force", "distance"]),
    ("1.7.2", ["kinetic energy", "speed of the motorbike", "has 81 j", "620 kj"]),
    ("1.7.3", ["gravitational potential", "mgh", "height", "stairs"]),
    ("1.7.4", ["conservation of energy", "air resistance", "transferred into the air", "energy of the penny"]),
    ("1.7.1", ["energy store", "energy stores", "chemical", "elastic", "magnetic", "kinetic"]),
    ("1.9.2", ["pressure due to", "fresh water", "sea water", "depth", "fish tank", "submarine"]),
    ("1.9.1", ["pressure", "force per unit area"]),
    ("1.2.1", ["speed", "velocity", "average velocity"]),
    ("1.2.2", ["acceleration", "accelerates", "constant acceleration", "resultant force acting"]),
    ("1.2.3", ["position-time", "distance-time", "position as a function"]),
    ("1.2.4", ["velocity-time", "speed-time", "speed time", "velocity as a function of time"]),
    ("1.2.5", ["average acceleration", "speed changes with time", "from speed-time"]),
    ("1.2.6", ["free fall", "freefall", "drops", "falls", "tosses", "apple straight up", "acceleration of free fall"]),
    ("1.3.1", ["mass and weight", "spring balance", "beam balance", "moon", "gravitational field strength"]),
    ("1.3.2", ["density", "mass balance", "cube", "g cm"]),
    ("1.4.1", ["resultant force", "largest magnitude of acceleration", "forces shown"]),
    ("1.4.3", ["newton", "f=ma", "force", "mass", "acceleration"]),
    ("1.6.1", ["momentum", "collision", "strikes the second", "velocity after", "rebounds"]),
    ("1.6.2", ["impulse", "contact with the ball", "ns"]),
]


def classify_question(question, lookup):
    text = question.lower()
    best, score, evidence = None, 0, []
    for code, keywords in RULES:
        s, hits = 0, []
        for kw in keywords:
            if kw in text:
                s += 3 if " " in kw else 1
                hits.append(kw)
        if s > score:
            best, score, evidence = code, s, hits
    if not best:
        for code, row in lookup.items():
            words = [w.lower() for w in re.findall(r"[A-Za-z]+", row["subtopic"]) if len(w) > 3]
            hits = [w for w in words if w in text]
            if len(hits) > score:
                best, score, evidence = code, len(hits), hits
    if best and best in lookup:
        row = lookup[best]
        confidence = "High" if score >= 3 else "Medium" if score > 0 else "Low"
        return row["topic"], row["chapter"], row["code"], row["subtopic"], confidence, ", ".join(evidence[:5])
    return "Unmapped", "Unmapped", "", "Unmapped", "Low", "No strong keyword match"


def analyze_files(qp_path, syllabus_path):
    qp_pages = extract_pdf_text(qp_path)
    syllabus_pages = extract_pdf_text(syllabus_path)
    syllabus_rows = parse_syllabus(syllabus_pages)
    lookup = {r["code"]: r for r in syllabus_rows}
    questions = parse_questions(qp_pages)
    records = []
    for q in questions:
        topic, chapter, code, subtopic, confidence, evidence = classify_question(q["question"], lookup)
        records.append({
            **q,
            "marks": 1,
            "topic": topic,
            "chapter": chapter,
            "subtopic_code": code,
            "subtopic": subtopic,
            "confidence": confidence,
            "evidence": evidence,
        })
    df = pd.DataFrame(records)
    return df, syllabus_rows


def summary_tables(df):
    total = int(df["marks"].sum()) if not df.empty else 0
    chapter = df.groupby("chapter", dropna=False).agg(questions=("question_no", "count"), marks=("marks", "sum")).reset_index()
    chapter["weightage_%"] = (chapter["marks"] / total * 100).round(1) if total else 0
    chapter = chapter.sort_values(["marks", "chapter"], ascending=[False, True])
    sub = df.groupby(["chapter", "subtopic_code", "subtopic"], dropna=False).agg(questions=("question_no", "count"), marks=("marks", "sum")).reset_index()
    sub["weightage_%"] = (sub["marks"] / total * 100).round(1) if total else 0
    sub = sub.sort_values(["marks", "chapter", "subtopic"], ascending=[False, True, True])
    return chapter, sub


def make_docx_report(df, output_path, title="Question Paper Topic Weightage Report"):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d-%b-%Y %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total = int(df["marks"].sum()) if not df.empty else 0
    doc.add_heading("Overall Summary", level=1)
    doc.add_paragraph(f"Total questions parsed: {len(df)}")
    doc.add_paragraph(f"Total marks analysed: {total}")
    chapter, sub = summary_tables(df)
    for heading, data, first_col in [
        ("Chapter-wise Weightage", chapter, "chapter"),
        ("Subtopic-wise Weightage", sub, "subtopic"),
    ]:
        doc.add_heading(heading, level=1)
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = first_col.title()
        hdr[1].text = "Questions"
        hdr[2].text = "Marks"
        hdr[3].text = "Weightage %"
        for _, r in data.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(r[first_col])
            cells[1].text = str(int(r["questions"]))
            cells[2].text = str(int(r["marks"]))
            cells[3].text = str(r["weightage_%"])
    doc.add_heading("Question-wise Topic Mapping", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Q.No", "Question", "Chapter", "Subtopic Code", "Subtopic", "Marks", "Confidence"]
    for i, head in enumerate(headers):
        table.rows[0].cells[i].text = head
    for _, r in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r["question_no"])
        cells[1].text = str(r["question"])[:700]
        cells[2].text = str(r["chapter"])
        cells[3].text = str(r["subtopic_code"])
        cells[4].text = str(r["subtopic"])
        cells[5].text = str(r["marks"])
        cells[6].text = str(r["confidence"])
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)
    doc.save(output_path)


def make_excel_report(df, output_path):
    chapter, sub = summary_tables(df)
    with pd.ExcelWriter(output_path, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Question Mapping", index=False)
        chapter.to_excel(writer, sheet_name="Chapter Weightage", index=False)
        sub.to_excel(writer, sheet_name="Subtopic Weightage", index=False)
        for sheet_name in writer.sheets:
            ws = writer.sheets[sheet_name]
            ws.set_column(0, 0, 18)
            ws.set_column(1, 1, 65)
            ws.set_column(2, 6, 22)
